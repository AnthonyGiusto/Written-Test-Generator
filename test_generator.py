"""
This application provides a tool to create a question database for a test
 and allows for the test administrator to automatically create multiple
 versions of the test in Microsoft with  randomized orders of the questions.
"""

from docx import Document
from secrets import choice
from random import randint
from random import shuffle


def question_gen():
    """ Allows the user to create a pool of test questions and create
    different versions of the same test"""

    question_bank = []
    initprompt = input(
        "Would you like to enter a test pool? Enter Y or N: ").lower()
    if initprompt == 'y':
        pool_name = input("Please enter the name of the test pool: ")
    else:
        exit()
    question_inc = 0

    # Creates a blank docx and saves the initial test-pool.
    original_doc = Document()
    original_doc.save(f"{pool_name}_original.docx")
    testpool = 'y'

    # Creates the question-pool and assigns it to the question_bank list.
    while testpool == 'y':
        question = input(
            f"Enter the text for question #{(question_inc)+1}: ")
        question_bank.append(question)
        question_inc = question_inc + 1
        original_doc.add_paragraph(f"{question} \n")
        original_doc.save(f"{pool_name}_original.docx")
        testpool = input("Would you like to continue? Enter Y or N: ").lower()

    # Generates the different versions of the test.
    version = input("How many versions would you like? (Enter up to 10): ")
    version = int(version)
    newdoc = Document()
    while version > 10:  # Keeps user from defining large number of docx files.
        print("You cannot create more than 10 versions.")
        version = input(
            "How many versions would you like? (Enter up to 10): ")

    # Creates a docx for each version and assigns the questions.
    question_num = 1
    while version:
        newdoc = Document()
        newdoc.save(f"{pool_name}_version_{version}.docx")
        shuffle(question_bank)

        # Writes each shuffled question to version docx
        for question in question_bank:
            newdoc.add_paragraph(f"{question_num}.) {question}\n\n")
        newdoc.save(f"{pool_name}_version_{version}.docx")
        version -= 1
        question_num += 1


question_gen()
